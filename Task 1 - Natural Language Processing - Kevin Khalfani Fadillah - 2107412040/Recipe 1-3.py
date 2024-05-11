# #Import library
# from docx import Document

# #Creating a word file object
# doc = open("file.docx","rb")
# #creating word reader object
# document = docx.Document(doc)
# # create an empty string and call this document. This document variable store each paragraph in the Word document.We then create a for loop that goes through each paragraph in the Word document and appends the paragraph.
# docu=""
# for para in document.paragraphs:
#  docu += para.text
# #to see the output call docu
# print(docu)

# Import library
from docx import Document

# Open the Word file and create a Document object
doc = Document("file.docx")

# Create an empty string to store the document text
doc_text = ""

# Iterate through each paragraph in the document and append the text to the doc_text variable
for para in doc.paragraphs:
    doc_text += para.text

# Print the document text
print(doc_text)
